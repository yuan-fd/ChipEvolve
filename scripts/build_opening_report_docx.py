#!/usr/bin/env python3
"""Build the submission DOCX from the reviewed HTML and rendered PNG figures."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import html

DESKTOP = Path("/share/home/yuanwenjie/Desktop")
SOURCE = DESKTOP / "AgenticEDA_开题报告_终稿.html"
FIGURES = DESKTOP / "AgenticEDA_开题报告_图片_终稿"
OUTPUT = DESKTOP / "AgenticEDA_开题报告_终稿.docx"


def set_cell_shading(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    props.append(shade)


def set_repeat_table_header(row) -> None:
    props = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    props.append(repeat)


def add_text_paragraph(document: Document, value: str, style: str | None = None,
                       *, bold_prefix: str | None = None):
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.35
    if bold_prefix and value.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        paragraph.add_run(value[len(bold_prefix):])
    else:
        paragraph.add_run(value)
    return paragraph


def add_table(document: Document, element) -> None:
    rows = element.xpath("./tr|./thead/tr|./tbody/tr")
    if not rows:
        return
    count = max(len(row.xpath("./th|./td")) for row in rows)
    table = document.add_table(rows=0, cols=count)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, source_row in enumerate(rows):
        row = table.add_row()
        for col_index, source_cell in enumerate(source_row.xpath("./th|./td")):
            cell = row.cells[col_index]
            cell.text = " ".join(source_cell.text_content().split())
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if source_cell.tag == "th" or row_index == 0:
                set_cell_shading(cell, "DDEBF7")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
        if row_index == 0:
            set_repeat_table_header(row)
    document.add_paragraph()


def add_figure(document: Document, element) -> None:
    image = element.xpath("./img")
    if not image:
        return
    src = image[0].get("src", "")
    prefix = Path(src).name.split("_", 1)[0]
    png = FIGURES / f"{prefix}_preview.png"
    if not png.is_file():
        raise FileNotFoundError(png)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(png), width=Cm(16.2))
    captions = element.xpath("./p[contains(@class,'caption')]")
    if captions:
        cap = document.add_paragraph(" ".join(captions[0].text_content().split()))
        cap.style = "Caption"
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build() -> None:
    root = html.fromstring(SOURCE.read_text(encoding="utf-8"))
    main = root.xpath("//main")[0]
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2); section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4); section.right_margin = Cm(2.2)

    styles = document.styles
    for name, size, bold, color in (
        ("Normal", 10.5, False, "1F2937"), ("Title", 22, True, "17375F"),
        ("Heading 1", 16, True, "1D4E7A"), ("Heading 2", 13, True, "285D86"),
        ("Caption", 9, False, "526175"),
    ):
        style = styles[name]
        style.font.name = "Arial"; style.font.size = Pt(size); style.font.bold = bold
        style.font.color.rgb = __import__("docx").shared.RGBColor.from_string(color)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for element in main:
        tag = element.tag.lower() if isinstance(element.tag, str) else ""
        value = " ".join(element.text_content().split())
        classes = set((element.get("class") or "").split())
        if not value and tag != "div":
            continue
        if tag == "h1":
            paragraph = document.add_paragraph(value, style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif tag == "h2":
            document.add_heading(value, level=1)
        elif tag == "h3":
            document.add_heading(value, level=2)
        elif tag == "p":
            style = "Subtitle" if "lead" in classes else None
            add_text_paragraph(document, value, style)
        elif tag in {"ol", "ul"}:
            style = "List Number" if tag == "ol" else "List Bullet"
            for item in element.xpath("./li"):
                add_text_paragraph(document, " ".join(item.text_content().split()), style)
        elif tag == "table":
            add_table(document, element)
        elif tag == "div" and "fig" in classes:
            add_figure(document, element)
        elif tag == "div" and (classes & {"box", "review", "formula"}):
            paragraph = add_text_paragraph(document, value)
            paragraph.paragraph_format.left_indent = Cm(.45)
            paragraph.paragraph_format.right_indent = Cm(.25)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(8)
        elif tag == "div":
            # Container-only elements are uncommon in the final report.  Add
            # their direct paragraphs/cards once instead of duplicating all
            # descendant text.
            for child in element.xpath("./p|./h3"):
                child_value = " ".join(child.text_content().split())
                if child_value:
                    add_text_paragraph(document, child_value,
                                       "Heading 2" if child.tag == "h3" else None)

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("基于 OpenROAD 的 AgenticEDA 自演化平台开题报告 · 2026")
    document.core_properties.title = "基于OpenROAD的AgenticEDA自演化平台开题报告"
    document.core_properties.subject = "AgenticEDA, OpenROAD, RTLScout, EDAIR, Bayesian Optimization"
    document.core_properties.author = "项目组"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
