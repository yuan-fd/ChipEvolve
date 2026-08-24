#!/usr/bin/env python3
"""Build the reviewed AgenticEDA submission DOCX from the frozen HTML."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import html


DESKTOP = Path("/share/home/yuanwenjie/Desktop")
SOURCE = DESKTOP / "AgenticEDA_开题报告_四轮审稿版.html"
FIGURES = DESKTOP / "AgenticEDA_开题报告_四轮审稿版_图片"
OUTPUT = DESKTOP / "AgenticEDA_开题报告_四轮审稿版.docx"


def chinese_font(style, name: str, size: float, *, bold=False, color="172033") -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def shade(cell, color: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), color)
    props.append(element)


def repeat_header(row) -> None:
    props = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    props.append(element)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText"); instruction.set(qn("xml:space"), "preserve"); instruction.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def paragraph(document: Document, value: str, style: str | None = None):
    p = document.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.35
    p.add_run(value)
    return p


def add_table(document: Document, source) -> None:
    rows = source.xpath("./tr|./thead/tr|./tbody/tr")
    if not rows:
        return
    columns = max(len(row.xpath("./th|./td")) for row in rows)
    table = document.add_table(rows=0, cols=columns)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, source_row in enumerate(rows):
        row = table.add_row()
        for col_index, source_cell in enumerate(source_row.xpath("./th|./td")):
            cell = row.cells[col_index]
            cell.text = " ".join(source_cell.text_content().split())
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(1.5)
                p.paragraph_format.line_spacing = 1.08
                for run in p.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(8.3 if columns >= 5 else 9)
            if source_cell.tag == "th" or row_index == 0:
                shade(cell, "E6EEF7")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
        if row_index == 0:
            repeat_header(row)
    document.add_paragraph()


def add_figure(document: Document, source) -> None:
    image = source.xpath("./img")
    if not image:
        return
    prefix = Path(image[0].get("src", "")).name.split("_", 1)[0]
    png = FIGURES / f"{prefix}_preview.png"
    if not png.is_file():
        raise FileNotFoundError(png)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(png), width=Cm(16.0))
    captions = source.xpath("./p[contains(@class,'caption')]")
    if captions:
        cap = document.add_paragraph(" ".join(captions[0].text_content().split()), style="Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build() -> None:
    root = html.fromstring(SOURCE.read_text(encoding="utf-8"))
    main = root.xpath("//main")[0]
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.3); section.right_margin = Cm(2.1)
    for style_name, size, bold, color in (
        ("Normal", 10.5, False, "172033"), ("Title", 22, True, "162A46"),
        ("Subtitle", 12, False, "40536B"), ("Heading 1", 16, True, "1C416E"),
        ("Heading 2", 13, True, "28577D"), ("Heading 3", 11, True, "354B66"),
        ("Caption", 8.7, False, "526074"),
    ):
        chinese_font(document.styles[style_name], "Microsoft YaHei", size, bold=bold, color=color)

    first_title = True
    for element in main:
        if not isinstance(element.tag, str):
            continue
        tag = element.tag.lower()
        classes = set((element.get("class") or "").split())
        value = " ".join(element.text_content().split())
        if tag == "h1":
            p = document.add_paragraph(value, style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            first_title = False
        elif tag == "h2":
            document.add_heading(value, level=1)
        elif tag == "h3":
            document.add_heading(value, level=2)
        elif tag == "h4":
            document.add_heading(value, level=3)
        elif tag == "p":
            p = paragraph(document, value, "Subtitle" if "lead" in classes else None)
            if "meta" in classes:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif tag in {"ol", "ul"}:
            style = "List Number" if tag == "ol" else "List Bullet"
            for item in element.xpath("./li"):
                paragraph(document, " ".join(item.text_content().split()), style)
        elif tag == "table":
            add_table(document, element)
        elif tag == "div" and "fig" in classes:
            add_figure(document, element)
        elif tag == "div" and classes.intersection({"statement", "plain", "warning", "evidence-box", "review", "formula"}):
            p = paragraph(document, value)
            p.paragraph_format.left_indent = Cm(0.45)
            p.paragraph_format.right_indent = Cm(0.25)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(8)

    footer = document.sections[0].footer.paragraphs[0]
    footer.add_run("基于OpenROAD的AgenticEDA自演化平台开题报告（四轮审稿版） · ")
    add_page_number(footer)
    document.core_properties.title = "基于OpenROAD的AgenticEDA自演化平台开题报告"
    document.core_properties.subject = "自然语言RTL、AgenticEDA、EDAIR、BO/GP、自演化知识"
    document.core_properties.author = "项目组"
    document.core_properties.comments = "事实核查与四轮审稿日期：2026-08-25"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
