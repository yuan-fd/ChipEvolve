#!/usr/bin/env python3
"""Generate factual v2 HTML/DOCX reports from frozen acceptance evidence."""

from __future__ import annotations

import html
import json
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path("/share/home/yuanwenjie/Desktop")
FIGURES = DESKTOP / "OpenROAD_AgenticEDA_v2_验收报告_图片"
MASTER_HTML = DESKTOP / "OpenROAD_AgenticEDA_v2_验收与开题报告_2026-08-25.html"
MASTER_DOCX = DESKTOP / "OpenROAD_AgenticEDA_v2_验收与开题报告_2026-08-25.docx"


def load(relative: str) -> dict[str, Any]:
    return json.loads((ROOT / relative).read_text(encoding="utf-8"))


RTL = load("artifacts/v2-real-rtl-suite-20260825/aggregate.json")
LOOP = load("artifacts/v2-multidesign-closed-loop-20260825/aggregate.json")
PARAM = load("artifacts/v2-parameter-ablation-multiseed-20260825/aggregate.json")
LEARNING = load("artifacts/v2-learning-ablation-20260825/aggregate.json")
EDAIR = load("artifacts/v2-edair-ablation-20260825/aggregate.json")
AGENT = load("artifacts/v2-agent-architecture-20260825/aggregate.json")
LEDGER = load("artifacts/v2-acceptance-20260825/manifest.json")
KNOWLEDGE = load("knowledge/public-corpus.lock.json")


def esc(value: Any) -> str:
    return html.escape(str(value), quote=True)


def table(headers: Iterable[str], rows: Iterable[Iterable[Any]], cls: str = "") -> str:
    head = "".join(f"<th>{esc(item)}</th>" for item in headers)
    body = "".join("<tr>" + "".join(f"<td>{item if isinstance(item, Safe) else esc(item)}</td>" for item in row) + "</tr>" for row in rows)
    return f'<table class="{cls}"><thead><tr>{head}</tr></thead><tbody>{body}</tbody></table>'


class Safe(str):
    pass


def tags(*items: str) -> Safe:
    return Safe(" ".join(f'<span class="tag">{esc(item)}</span>' for item in items))


def figure(number: int, caption: str) -> str:
    name = f"{number:02d}_drawio_style.svg"
    return (f'<figure><img src="OpenROAD_AgenticEDA_v2_验收报告_图片/{name}" '
            f'alt="Figure {number}"><figcaption>图 {number}　{esc(caption)}</figcaption></figure>')


PAPERS = {
    "rtl": [
        ("EDA-Aware RTL Generation with Large Language Models", "DATE 2025", "10.23919/DATE64628.2025.10992789", "把 EDA 日志反馈带入 RTL 迭代，而不是一次生成即结束。", "已对齐工具反馈、自动修订和独立验证包；尚无多 seed 泛化统计。"),
        ("VeriOpt: PPA-Aware High-Quality Verilog Generation via Multi-Role LLMs", "ICCAD 2025", "10.1109/ICCAD66269.2025.11240771", "Planner/Programmer/Reviewer/Evaluator 多角色并使用 PPA 反馈。", "已有 Spec/Verification/RTL/Runtime 角色和后端 PPA；尚未证明多角色本身提高 QoR。"),
        ("AutoSilicon: Scaling Up RTL Design Generation Capability of Large Language Models", "ACM TODAES 2025", "10.1145/3737286", "将复杂 RTL 任务分解成较小模块并即时 compile/test。", "当前固定题库验证了中小模块；层级 SpecIR 和多模块集成回归仍是扩展项。"),
        ("EvolVE: Evolutionary Search for LLM-based Verilog Generation and Optimization", "arXiv:2601.18067v1", "2601.18067v1", "使用 evolutionary/MCTS 式候选搜索和结构化测试生成。", "当前有有界候选迭代和 lineage；未实现论文的完整搜索算法。"),
    ],
    "agent": [
        ("ORFS-agent: Tool-Using Agents for Chip Design Optimization", "MLCAD 2025", "10.1109/MLCAD65511.2025.11189204", "Agent 调用 ORFS 工具并依据 QoR 反馈迭代。", "已有 Runtime 工具边界、BO/GP 闭环和八阶段 trace；方法不等同于复现 ORFS-agent。"),
        ("AgenticPD: A Stage-Aware Agentic Framework for Physical Design QoR Optimization", "arXiv:2607.04758v2", "2607.04758v2", "Judge、stage agents、结构化历史和 checkpoint branching。", "已有 checkpoint、阶段诊断和恢复测试；v2 不含各阶段 repair 执行器。"),
        ("From Tool Invocation to Source-Mechanism Exploration: Protected White-Box DSE for Open-Source EDA", "arXiv:2607.11294v4", "2607.11294v4", "Teacher/Students、私有 workspace、受保护 evaluator 和机制证据。", "已有不可执行 hypothesis、Runtime 唯一执行权和 holdout；白盒源码探索属于后续工具阶段。"),
        ("PDAGENT-BENCH", "arXiv:2606.17253v6", "2606.17253v6", "以工具执行和长程物理设计任务评估 Agent。", "当前用真实四设计 trace 和故障注入测编排；尚无 353 题规模 benchmark。"),
    ],
    "data": [
        ("CircuitOps: An ML Infrastructure Enabling Generative AI for VLSI Circuit Optimization", "ICCAD 2023", "10.1109/ICCAD57390.2023.10323611", "把 cell/pin/net 及其关系转换成适合 ML 的表和图。", "当前从综合网表导出逻辑关系，并与 EDAIR、artifact SHA 和 loss manifest 联用。"),
        ("EDATracer: An Agentic Framework for Large-Scale EDA Artifact Analysis", "arXiv:2608.04032v1", "2608.04032v1", "跨 source/scripts/log/netlist/report 建知识图并做语义索引。", "当前有 typed IR、关系表和分级回读；没有其大规模知识图与向量双索引。"),
        ("Customized Retrieval Augmented Generation and Benchmarking for EDA Tool Documentation QA", "ICCAD 2024", "10.1145/3676536.3676730", "针对 EDA 文档问答定制 RAG 与 benchmark。", "知识库保存精确文献元数据和受限 claim；当前检索是 context 过滤 + BM25，不是向量 RAG。"),
    ],
    "bo": [
        ("PTPT: Physical Design Tool Parameter Tuning via Multi-Objective Bayesian Optimization", "IEEE TCAD 2023", "10.1109/TCAD.2022.3167858", "以多目标 BO 调物理设计工具参数。", "当前是透明的每目标固定 RBF GP + 加权 EI，不是 Pareto MOBO。"),
        ("HyperPlace: Harnessing a Large Language Model for Efficient Hyperparameter Optimization in GPU-Accelerated VLSI Placement", "ACM TODAES 2025", "10.1145/3733601", "使用 LLM in-context batch HPO 优化 GPU placement。", "仅作为批量 HPO 对照；本平台当前算法是 GP，不声称复现。"),
        ("iPO: Constant Liar Parameter Optimization for Placement with Representation and Transfer Learning", "ACM TODAES 2026", "10.1145/3747292", "并行 Constant-Liar 优化及 representation/transfer learning。", "当前没有 Constant-Liar 和跨设计 surrogate transfer；已有经验迁移的独立 holdout 门。"),
    ],
    "learning": [
        ("Retrieve, Schedule, Reflect: LLM Agents for Chip QoR Optimization", "arXiv:2603.13767v2", "2603.13767v2", "检索历史、调度实验并依据 Pareto 反馈反思。", "已有 evidence retrieval、实验 trace 和反思 hypothesis；当前检索不是向量模型。"),
        ("ReviewDSE", "arXiv:2607.11294v4", "2607.11294v4", "保存机制级正负证据，并用受保护 evaluator 审查。", "已有正负知识、2×2 干预和 held-out admission；仅有一个真实迁移反例。"),
        ("AgenticPD", "arXiv:2607.04758v2", "2607.04758v2", "使用阶段历史与结构化决策记录。", "已有 context fingerprint、八阶段 trace 和 append-only hypothesis ledger。"),
    ],
}


def paper_table(key: str) -> str:
    return table(("论文/项目", "来源", "它解决什么", "本平台对齐与差距"),
                 ((f"{title}（{identifier}）", venue, method, mapping)
                  for title, venue, identifier, method, mapping in PAPERS[key]))


def rtl_section() -> str:
    rows = []
    for row in sorted(RTL["design_rows"], key=lambda x: x["design"]):
        m = row["mutation"]
        rows.append((row["design"], row["candidate_files"], row["rtl_revisions"],
                     f'{m["killed_count"]}/{m["executable_count"]}',
                     f'{m["generated_count"]} 生成、{m["invalid_count"]} 无效',
                     f'{m["mutation_score"]:.2f}', row["gds"]["size_bytes"]))
    return f"""
<section id="rtl"><h2>3. RTL 生成与验证：自动化，但写 RTL 与写测试分开</h2>
<div class="plain"><b>大白话结论：</b>用户输入自然语言，平台先形成 SpecIR；Verification Agent 自动写并冻结 testbench；RTLScout 只能写 RTL，不能改判题规则；固定工具链做 lint、综合、仿真和 mutation；通过后再进入 OpenROAD。所谓“人工审核人”字段只是证据归属，不表示默认要人手写测试。</div>
{figure(4, "RTLScout 被拆成可观察的候选生成、固定评估、反馈修补和择优循环。")}
<h3>3.1 当前算法链</h3>
<ol><li>Spec Agent 将功能、端口、位宽、时钟/复位、约束、验收条件和未决问题写进 SpecIR。</li>
<li>Verification Agent 在独立上下文中生成自检 testbench、假设和覆盖计划；结构预检确保有 DUT 实例、时钟推进、失败路径和机器可读汇总。</li>
<li>testbench 以 SHA-256 冻结成 VerificationPackage。RTLScout 没有修改它的权限。</li>
<li>RTLScout 在隔离 workspace 内执行“计划→受限文件编辑→compile/lint→仿真→读 diagnostic→最小修补/分支”。</li>
<li>mutation gate 故意制造一组小错误，看 testbench 能杀死多少个；它用于发现“测试太弱”，不是形式化证明。</li>
<li>只有正确性门通过的 RTL 才按 Yosys/ABC cost 排序；最终 PPA 必须由后端 OpenROAD 重测。</li></ol>
<h3>3.2 真实固定题库结果</h3>{table(("设计", "候选文件", "自动修订", "杀死/可执行 mutant", "生成/无效 mutant", "mutation score", "GDS 字节"), rows)}
<div class="boundary"><b>证据边界：</b>{esc(RTL['claim_boundary'])} GCD 有 8/16 个 mutant 可执行，UART 有 6/15 个可执行，不能只展示 1.0 分数而隐去无效 mutant；ibex_alu 只是 Ibex 风格 ALU 小模块，不是完整 Ibex 核。</div>
<h3>3.3 与前沿逐项比较</h3>{paper_table('rtl')}</section>"""


def agent_section() -> str:
    rows = []
    for row in sorted(AGENT["design_rows"], key=lambda x: x["design"]):
        rows.append((row["design"], row["status"], row["event_count"],
                     " / ".join(f"{k}:{v}" for k, v in row["phase_counts"].items())))
    return f"""
<section id="agent"><h2>4. Agent 架构：一套状态协议，不是多个聊天框</h2>
<div class="plain"><b>大白话结论：</b>Agent 可以理解、提假设和安排实验，但没有直接执行权；Runtime 才能启动 EDA 工具、登记结果和产物。八阶段协议让每一步“看了什么、为什么提案、谁验证、学到了什么”都可追踪。</div>
<h3>4.1 八阶段输入输出</h3>{table(("阶段", "必须产出", "权限边界"), [
('map 地图','设计/工艺/工具/阶段/历史运行地图','不能执行命令'),('semantic 语义','规格语义或结构化诊断，显式 unknown','不能编造未解析字段'),('experiment 实验','变量、对照、重复、预算、停止规则','不能修改服务器策略'),('hypothesis 假设','机制、依据、预期和 falsifier','永远不可执行'),('implement 实现','Runtime 可消费的参数向量/ActionSpec','执行权仍属于 Runtime'),('validate 验证','真实 run_id、metric、artifact SHA','不能使用 Agent 自报分数'),('review 审查','硬约束、统计改善、接受/拒绝','不能只看最好单点'),('memory 记忆','正/负知识卡和适用边界','未经 holdout 不得扩展范围')])}
<h3>4.2 真实四设计 trace</h3>{table(("设计", "终态", "事件数", "阶段计数"), rows)}
<p>另外 14 项边界测试覆盖 baseline/candidate 中断恢复、proposal identity 保持、终态幂等、clock 不得进入搜索空间、浏览器不得注入 seed/repetition/transition/search 控制。</p>
<div class="boundary"><b>证据边界：</b>{esc(AGENT['claim_boundary'])}</div>
<h3>4.3 与前沿逐项比较</h3>{paper_table('agent')}</section>"""


def edair_section() -> str:
    t = EDAIR["totals"]
    return f"""
<section id="edair"><h2>5. EDA-to-AI 数据接口：摘要、对象和原文三层并存</h2>
<div class="plain"><b>大白话结论：</b>AI 不是只读 12 个 KPI，也不是把几百 MB 日志全部塞进 prompt。平台保留原文件作权威，再生成 timing path、instance、net 等结构化对象，最后按问题抽取一个有上限的 evidence packet；被省略的内容写进 loss manifest，并可凭 artifact ID 回读。</div>
{figure(3, "原始 artifact、版本化 EDAIR/CircuitOps 和有界 Agent 查询层同时存在。")}
<h3>5.1 四设计导出规模</h3>{table(("项目", "数量", "含义"), [
('逻辑 instance',t['logical_instance_objects'],'从 registered netlist 恢复的逻辑对象'),('逻辑 net',t['logical_net_objects'],'可查询连接关系'),('物理 instance',t['physical_instance_objects'],'从 DEF 恢复的位置对象'),('OpenSTA timing path',t['timing_path_objects'],'带 path/point 结构的时序证据'),('raw artifact 目录',t['raw_artifact_directory'],'每项带 SHA-256 和来源'),('Agent facts',t['agent_packet_facts'],'按诊断目标有界抽取')])}
<h3>5.2 fidelity 机制</h3><ul><li>每个对象携带 parser 和原始 artifact 指针；解析不到时保持 unknown，不填假值。</li><li>逻辑网表关系与物理 DEF 实例分开命名，避免把“逻辑连接”冒充“物理布线”。</li><li>loss manifest 明确 packet 未内联的 path、net、violation 和原始文件。</li><li>AI 可先读高层事实，不足时按 artifact ID、offset 和长度回读原文。</li></ul>
<div class="boundary"><b>证据边界：</b>{esc(EDAIR['claim_boundary'])} 当前还不是 EDATracer 规模的知识图+向量双索引，也不是完整 ODB/Liberty 属性导出。</div>
<h3>5.3 与前沿逐项比较</h3>{paper_table('data')}</section>"""


def bo_section() -> str:
    rows = []
    for row in PARAM["design_rows"]:
        rows.append((row["design"], f'{row["bo_best_utility"]["median"]*100:.3f}%',
                     f'{row["random_best_utility"]["median"]*100:.3f}%', row["median_winner"],
                     f'{row["bo_threshold_seed_rate"]*3:.0f}/3', f'{row["random_threshold_seed_rate"]*3:.0f}/3'))
    return f"""
<section id="bo"><h2>6. BO/GP 参数探索：唯一产品优化入口</h2>
<div class="plain"><b>大白话结论：</b>用户不再选择 baseline、顺序扫描或 Agent 三种模式。按下一个按钮后，服务器先把 baseline 重复测三次，再让 BO/GP 挑“下一组最值得真实测的组合”。连续三轮没有达到改善阈值，就停止调参并转入阶段诊断。</div>
{figure(6, "当前 BO 的真实算法、重复实验合同、停滞换向和等预算结果。")}
<h3>6.1 当前代码具体算什么</h3><ol><li>搜索变量当前是连续参数；每维按上下界映射到 [0,1]，clock 被硬性冻结，不能靠放宽时钟伪造改善。</li><li>同一参数组合的重复运行先聚合均值，并用 sample variance / replica 数作为该均值的 observation noise。</li><li>area、timing、power 每个目标分别拟合一个 exact RBF GP，固定 length scale=0.35，以 Cholesky 求解。</li><li>每个目标按方向和观测 min/max 归一化，再按 balanced/timing/area/power profile 的权重合成 utility。</li><li>计算 scalarized Expected Improvement；失败点不训练 QoR GP，但通过邻域 Beta(1,1) 平滑经验可行率降低附近 acquisition。</li><li>每轮生成 512 点 deterministic Latin-hypercube 候选池，剔除已测点，选择 acquisition 最大的组合。</li><li>GP 输出只是 ExperimentPlan；OpenROAD 完整 finish 的重复实测才有资格更新 best utility。</li></ol>
<h3>6.2 等预算三 seed 结果</h3>{table(("设计", "BO 中位最佳 utility", "Random 中位最佳 utility", "中位胜者", "BO 过 0.5%", "Random 过 0.5%"), rows)}
<p>总预算：BO {PARAM['bo_run_count']} 次、Random {PARAM['random_run_count']} 次完整 flow。阈值命中 design-seed 单元为 BO {PARAM['bo_threshold_events']}/12，Random {PARAM['random_threshold_events']}/12；设计级中位数胜负为 {PARAM['median_design_wins']['bo']}:{PARAM['median_design_wins']['random']}。</p>
<div class="boundary"><b>允许的结论：</b>在这组固定设计、三 seed 和相同预算下，BO 提高了达到预注册 0.5% 实用阈值的频次。<b>不允许的结论：</b>BO 普遍优于随机、统计显著，或当前已实现 Pareto MOBO/qNEHVI。</div>
<h3>6.3 与前沿逐项比较</h3>{paper_table('bo')}</section>"""


def learning_section() -> str:
    arm_rows = [(row["arm"], row["stores_observations"], row["creates_transfer_rule"],
                 row["false_transfer_rule"], row["result"]) for row in LEARNING["arms"]]
    return f"""
<section id="learning"><h2>7. 自演化学习：不是记录成败，而是让经验经得起反例</h2>
<div class="plain"><b>大白话结论：</b>“上次把密度调高成功了”不是可迁移知识。平台先固定上下文和 RTL 指纹，做重复 2×2 组合干预，分开主效应与交互；再把同一假设预注册到没见过的设计上复测。方向反转就保存为 negative-transfer 证据，明确禁止自动复用。</div>
{figure(2, "运行事实经过局部干预和 held-out 复验，才可能进入知识生命周期。")}
{figure(5, "知识卡保存上下文、统计、机制、反证条件、run/artifact 指针和准入状态。")}
<h3>7.1 真实 GCD→FIFO 因果门实验</h3><p>GCD 的 utilization×density 面积交互为 <b>{LEARNING['source_interaction']:.3f} µm²</b>；预注册 FIFO holdout 后得到 <b>{LEARNING['holdout_interaction']:.3f} µm²</b>，方向反转。因此 ledger 从 draft→assessment→refuted 追加记录，`action_eligible=false`。24/24 次 full-flow 成功，WNS≥0 且 DRC=0。</p>
{table(("消融 arm", "保存观测", "形成迁移规则", "错误迁移规则", "结果"), arm_rows)}
<div class="boundary"><b>证据边界：</b>{esc(LEARNING['claim_boundary'])} “RAG-only 会错迁移”是对同一批真实数据构造的政策反事实，不是另一次线上 Agent 事故；当前不能声称达到普适因果推理或 2026 SOTA。</div>
<h3>7.2 知识库当前到底有什么</h3><p>当前知识锁包含 {len(KNOWLEDGE['sources'])} 个 source metadata、{len(KNOWLEDGE['claims'])} 条受限 claim 和 {len(KNOWLEDGE['benchmarks'])} 个 benchmark metadata。作者、年份、venue、DOI/arXiv 版本已进入校验契约；多数条目只保存元数据和受限方法陈述，不代表已经下载论文全文。检索版本是 <code>{esc(KNOWLEDGE['snapshot']['embedding_version'])}</code>，即 context 硬过滤 + BM25 风格词项检索，没有向量 embedding。</p>
<h3>7.3 与前沿逐项比较</h3>{paper_table('learning')}</section>"""


def experiments_section() -> str:
    records = []
    for row in LEDGER["evidence_records"]:
        records.append((row["capability"], row["status"], row["evidence_path"],
                        row["evidence_sha256"], row["claim_boundary"]))
    return f"""
<section id="experiments"><h2>8. 论文级实验清单与验收总账</h2>
<p>下面的 SHA-256 是 aggregate 文件本身的哈希；它用于证明报告引用的是哪份冻结证据。不同 aggregate 可能复用底层 run_id，因此不能把 48、288、24 等数字直接相加宣传“总运行次数”。</p>
{table(("模块", "状态", "权威 aggregate", "SHA-256", "结论边界"), records, 'evidence-table')}
<h3>8.1 已通过工程体检</h3><ul><li>完整仓库：327 passed，1 deselected。</li><li>Node 前端语法、Python compileall、JSON 解析和 git diff whitespace 均通过。</li><li>tracked secret scan 扫描 403 个文件，credential finding=0。</li><li>产品边界测试确认旧 baseline/顺序扫描/manual campaign/BYOK 路由不存在，浏览器不能注入 optimizer seed、重复数、transition 或搜索范围。</li></ul>
<h3>8.2 若投稿还应补的统计实验</h3><ul><li>RTL：四题至少 5 generation seeds，并完成 direct / RTLScout / independent verifier 三臂比较。</li><li>BO：增加设计、严格配对 policy seed，并报告 bootstrap interval 或分层模型；不能只扩大 run 数却不扩大独立设计单元。</li><li>EDAIR：建立人工标注诊断题，比较 KPI-only、raw log、EDAIR、EDAIR+关系图的定位准确率、证据率、token 和延迟。</li><li>Agent：做 flat orchestrator、无 checkpoint、无 review、完整八阶段的等预算消融，区分“编排可靠”与“QoR 提升”。</li><li>学习：增加第三个及更多 holdout 设计，报告负迁移率和规则校准，而不只展示一个反例。</li></ul></section>"""


def workflow_and_scope() -> str:
    return f"""
<section id="scope"><h2>1. 平台定位、产品边界与唯一入口</h2>
<div class="plain"><b>一句话定位：</b>这是 OpenROAD 上方的 AgenticEDA 实验控制与证据学习层，不是聊天式 tutorial。模型负责理解和提案；Runtime 负责执行；Verilator、Yosys、OpenSTA 和 OpenROAD 负责判定。</div>
{table(("用户可输入", "服务器自动负责", "用户不能手工选择"), [
('design_id、clock、platform、目标偏好','baseline 第0轮、BO/GP参数组合、3次重复、3轮停滞、诊断和证据归档','baseline模式、顺序扫描、grid/manual optimizer、seed/预算/范围'),
('自然语言设计说明','SpecIR、自动testbench、RTLScout候选、固定验证门','模型Provider、API Key、自带testbench绕过自动产品路径')])}
<p>正式写入口只有 <code>POST /api/v2/closed-loops</code> 和服务器托管的 <code>run-to-boundary</code>。baseline 没有被删除，因为科学比较必须有对照；它被降为 Agent 内部第 0 轮，不再是用户模式。统一模型是平台托管 <code>gpt-5.6-terra</code>；RTLScout 执行插件中的 Anthropic、DeepInfra、OpenRouter 及 API-Key 注入逻辑已删除。</p>
<div class="boundary"><b>研究边界：</b>v2 验收做到参数级自主闭环和诊断换向；TimingECO、Resynth、EvoDRC、DPLEvolve/GoalEvolve 源码动作执行器属于下一阶段，不用空壳提前占主链。</div></section>
<section id="workflow"><h2>2. 全流程 framework 与 workflow</h2>{figure(1, "最终目标总图；当前 v2 已跑通至诊断/学习边界，有限 repair 工具执行器在后续接入。")}
<p>主链是：自然语言→SpecIR→独立 Verification Agent→冻结 testbench→RTLScout 迭代→lint/simulation/mutation→OpenROAD baseline→重复 BO/GP→连续三轮停滞→阶段诊断→证据反思与知识准入。图中的 Repair Agent 已能生成诊断和受限动作提案，但 v2 不宣称已经执行 TimingECO/EvoDRC 等 v3 工具。</p></section>"""


def team_and_plan() -> str:
    return f"""
<section id="team"><h2>9. 课题分工、工作包与前沿对标</h2>
{table(("负责人", "工作包", "具体交付", "验收方式", "主要对标"), [
('袁文杰','平台主线、Agent、EDAIR、自演化、BO/GP、实验','唯一闭环API、八阶段trace、checkpoint、EDAIR/CircuitOps、知识准入、消融与报告','全链回放、权限/恢复测试、真实多设计对照','ORFS-agent 2025；AgenticPD/ReviewDSE/EDATracer 2026；PTPT/iPO'),
('刘宏博','RTL生成与验证','层级SpecIR、Verification Agent、RTLScout候选策略、mutation/coverage/formal门、固定与开放suite','多seed pass@k、mutation有效率、失败分类、同协议PPA分布','EDA-Aware RTL 2025；VeriOpt/AutoSilicon 2025；EvolVE 2026'),
('沈夏南','TaiWei3D维护与3D Agent','固定commit/工艺矩阵、cross-tier数据、3D阶段诊断与ECO门','官方case×工艺矩阵、2D/3D对照、fresh timing/DRC/连通性','AgenticECO 2026；3D partition/placement相关工作'),
('戴家乐','有限动作工具接口','TimingECO、物理感知Resynth、EvoDRC、源码演化器的manifest/ActionSpec/sandbox/rollback','每工具单独真实验收，随后跨阶段组合full-flow','EvoDRC/DRC-Aid/AgenticECO/ReviewDSE 2026；ReSynthAI 2025')])}
<h3>9.1 阶段边界</h3><ul><li><b>当前主线：</b>自然语言到 RTL、OpenROAD baseline、BO/GP、停滞诊断、EDAIR 和证据学习。</li><li><b>下一工作包：</b>逐个接入有限修复工具；每个工具先证明输入输出、白名单、回滚和 fresh full-flow verifier。</li><li><b>生态工作包：</b>验证多个工具是否能按阶段组合，不因局部优化破坏后续 timing/DRC/QoR。</li></ul></section>
<section id="risk"><h2>10. 风险、资源与商业合作验收</h2>
{table(("风险", "会怎样", "平台现有防线", "合作验收"), [
('自动testbench太弱','错误RTL也通过','写/审角色分离、冻结hash、mutation gate','扩大可执行mutant、coverage与reference/formal题'),
('BO样本少或模型偏','只在个别设计有效','重复统计、失败可行率、等预算random、结论边界','增加独立设计/seed并报告置信区间'),
('EDA转换丢细节','Agent误诊','原始artifact权威、SHA、loss manifest、按需回读','人工标注诊断benchmark和字段round-trip'),
('经验负迁移','旧规则伤害新设计','2×2干预、预注册holdout、negative evidence','更多leave-one-design-out和action eligibility审计'),
('外部工具破坏主链','局部修好、全流程变差','Runtime唯一执行、私有workspace、fresh verifier、rollback','每工具单独验收后做组合顺序消融'),
('私有PDK/设计泄漏','商业数据风险','租户隔离、artifact授权、模型无shell、secret扫描','部署侧访问控制、数据保留和IP协议')])}
<p>合作方需明确提供：可使用的设计/PDK/许可证、算力配额、允许发表的数据范围、第三方工具 license 和成果知识产权边界。建议首个 90 天以“扩充独立设计单元、完成 EDAIR 诊断 benchmark、接入一个有限动作工具”为验收，不以单次最好 PPA 作里程碑。</p></section>"""


def references() -> str:
    entries = []
    seen = set()
    for group in PAPERS.values():
        for title, venue, identifier, _, _ in group:
            key = (title, identifier)
            if key not in seen:
                entries.append(f"{title}. {venue}. {identifier}.")
                seen.add(key)
    return "<section id='references'><h2>11. 参考文献</h2><ol class='refs'>" + "".join(f"<li>{esc(x)}</li>" for x in entries) + "</ol><p class='small'>引用的 2026 arXiv 工作按所列版本视为预印本；其外部实验数字不是本平台结果。精确作者、venue、年份、DOI/arXiv 版本保存在 knowledge/public-corpus.lock.json。</p></section>"


CSS = """
:root{--ink:#172033;--muted:#5d6b7d;--blue:#315f9a;--line:#d8e0ea;--pale:#f4f7fb;--green:#27745b;--red:#9f3f49}*{box-sizing:border-box}body{margin:0;background:#edf1f5;color:var(--ink);font:16px/1.72 Arial,"Microsoft YaHei",sans-serif}main{max-width:1220px;margin:auto;background:white;padding:56px 70px 90px}h1{font-size:38px;line-height:1.25;margin:0 0 12px;color:#17365f}h2{font-size:25px;margin:52px 0 18px;padding:10px 15px;border-left:6px solid var(--blue);background:#f1f5fb}h3{font-size:19px;color:#244e7b;margin:30px 0 12px}p{margin:10px 0}.lead{font-size:20px;color:#3e5067}.meta{color:var(--muted)}nav{position:sticky;top:0;z-index:5;background:#fff;border:1px solid var(--line);padding:10px 14px;margin:25px 0}nav a{color:#285b91;margin-right:14px;text-decoration:none;white-space:nowrap}.hero{padding:24px;background:linear-gradient(135deg,#eef4ff,#f7fbf9);border:1px solid #cbd9ea;border-radius:14px}.metrics{display:grid;grid-template-columns:repeat(4,1fr);gap:12px;margin-top:18px}.metric{padding:15px;border:1px solid var(--line);border-radius:10px;background:white}.metric b{display:block;font-size:25px;color:#234f80}.plain,.boundary{padding:16px 18px;border-radius:9px;margin:15px 0}.plain{background:#eef6ff;border:1px solid #a9c2df}.boundary{background:#fff2f3;border:1px solid #d49aa1;color:#542c32}table{width:100%;border-collapse:collapse;margin:14px 0 25px;font-size:14px}th,td{border:1px solid #cfd8e4;padding:9px 10px;vertical-align:top}th{background:#eaf0f8;text-align:left;color:#24486e}tr:nth-child(even) td{background:#fbfcfe}figure{margin:26px -30px;padding:12px;border:1px solid var(--line);background:white}figure img{width:100%;display:block}figcaption{font-size:13px;color:var(--muted);padding:9px 6px 2px}.tag{display:inline-block;padding:2px 8px;border-radius:12px;background:#e8f1fb;color:#275c91;margin:2px}.refs,.small{font-size:13px;color:#4c596a}.evidence-table td:nth-child(4){font-family:monospace;word-break:break-all;font-size:11px}code{background:#edf1f5;padding:2px 5px;border-radius:4px}li{margin:6px 0}@media(max-width:800px){main{padding:28px 18px}.metrics{grid-template-columns:1fr 1fr}figure{margin:20px 0}table{font-size:12px}}@media print{body{background:white}main{max-width:none;padding:0 8mm}nav{display:none}h2{break-before:page}figure{margin:12px 0;break-inside:avoid}table{break-inside:auto}tr{break-inside:avoid}.hero{break-inside:avoid}}
"""


def master_html() -> str:
    metrics = f"""<div class="metrics"><div class="metric"><b>4/4</b>固定自然语言 RTL 到 GDS</div><div class="metric"><b>48</b>四设计自主闭环 full-flow</div><div class="metric"><b>7/12 vs 4/12</b>BO/Random 阈值事件</div><div class="metric"><b>327</b>完整测试通过</div></div>"""
    body = (workflow_and_scope() + rtl_section() + agent_section() + edair_section()
            + bo_section() + learning_section() + experiments_section()
            + team_and_plan() + references())
    return f"""<!doctype html><html lang="zh-CN"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>OpenROAD AgenticEDA v2 验收与开题报告</title><style>{CSS}</style></head><body><main>
<header class="hero"><p class="meta">技术能力、实验事实与合作实施方案 · 证据快照 2026-08-25</p><h1>OpenROAD AgenticEDA 自演化平台</h1><p class="lead">从自然语言 RTL 到可审计 OpenROAD 参数优化闭环</p><p>平台主线：袁文杰　·　TaiWei3D：沈夏南　·　RTL：刘宏博　·　有限动作工具：戴家乐</p>{metrics}<div class="boundary"><b>读者须知：</b>本报告把“代码存在”“真实实验已通过”“统计结论成立”“未来工具规划”分开。四设计 RTL 是单 seed 可行性；BO 结果是描述性对照；因果学习只拦住一次已观察的错误迁移。没有把这些事实包装成任意规格、普遍显著优势或 SOTA。</div></header>
<nav><a href="#scope">定位</a><a href="#workflow">流程</a><a href="#rtl">RTL</a><a href="#agent">Agent</a><a href="#edair">EDAIR</a><a href="#bo">BO/GP</a><a href="#learning">自演化</a><a href="#experiments">实验</a><a href="#team">分工</a><a href="#risk">风险</a></nav>{body}
<footer class="small"><hr>总验收账：artifacts/v2-acceptance-20260825/manifest.json · base commit {esc(LEDGER['base_git_commit'])} · 生成时工作树含本轮未提交改动，最终发布须以提交后的 commit 重新冻结。</footer></main></body></html>"""


def module_html(title: str, section: str) -> str:
    return f"<!doctype html><html lang='zh-CN'><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'><title>{esc(title)}</title><style>{CSS}</style></head><body><main><header class='hero'><h1>{esc(title)}</h1><p class='lead'>OpenROAD AgenticEDA v2 · 2026-08-25 真实证据版</p><p>本附件沿用总报告的 artifact、数字和 claim boundary，不使用旧报告快照。</p></header>{section}</main></body></html>"


def set_cell(cell, text: str, header: bool = False) -> None:
    cell.text = str(text)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = "Arial"; run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(8.5); run.bold = header
    if header:
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "DDEBF7"); cell._tc.get_or_add_tcPr().append(shd)


def add_doc_table(doc: Document, headers: tuple[str, ...], rows: list[tuple[Any, ...]]) -> None:
    tab = doc.add_table(rows=1, cols=len(headers)); tab.style = "Table Grid"
    for i, value in enumerate(headers): set_cell(tab.rows[0].cells[i], value, True)
    trpr = tab.rows[0]._tr.get_or_add_trPr(); repeat = OxmlElement("w:tblHeader"); repeat.set(qn("w:val"), "true"); trpr.append(repeat)
    for row in rows:
        cells = tab.add_row().cells
        for i, value in enumerate(row): set_cell(cells[i], value)
    doc.add_paragraph()


def add_doc_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style); p.paragraph_format.line_spacing = 1.35; p.paragraph_format.space_after = Pt(5)
    p.add_run(text)


def add_doc_figure(doc: Document, number: int, caption: str) -> None:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIGURES / f"{number:02d}_preview.png"), width=Cm(16.2))
    cap = doc.add_paragraph(f"图 {number}　{caption}", style="Caption"); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_docx() -> None:
    doc = Document(); sec = doc.sections[0]
    sec.top_margin=Cm(2); sec.bottom_margin=Cm(1.8); sec.left_margin=Cm(2.2); sec.right_margin=Cm(2.0)
    for name, size, color in (("Normal",10.5,"1F2937"),("Title",22,"17365F"),("Heading 1",16,"244E7B"),("Heading 2",13,"315F9A"),("Caption",9,"5D6B7D")):
        st=doc.styles[name]; st.font.name="Arial"; st._element.rPr.rFonts.set(qn("w:eastAsia"),"Microsoft YaHei"); st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color)
    title=doc.add_paragraph("OpenROAD AgenticEDA 自演化平台",style="Title"); title.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sub=doc.add_paragraph("v2 技术验收、开题与商业合作报告\n证据快照：2026-08-25"); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_doc_paragraph(doc,"负责人：袁文杰（平台）｜沈夏南（TaiWei3D）｜刘宏博（RTL）｜戴家乐（有限动作工具）")
    add_doc_paragraph(doc,"事实边界：四设计 RTL 为单 seed 固定题库可行性；BO 为描述性等预算对照；因果门拦住一次错误迁移。不得改写为任意规格、普遍显著优势或 SOTA。")
    doc.add_page_break()
    doc.add_heading("1. 平台定位与唯一入口",1)
    add_doc_paragraph(doc,"平台是 OpenROAD 上方的实验控制和证据学习层。用户只启动 Agent 自主 BO/GP 闭环；baseline 是服务器内部第0轮。顺序扫描、手工优化模式和浏览器 BYOK 已删除。模型固定为平台托管 gpt-5.6-terra。")
    add_doc_figure(doc,1,"从自然语言到 OpenROAD 搜索、停滞诊断与知识准入的总体闭环。")
    doc.add_heading("2. RTL 生成与验证",1)
    add_doc_paragraph(doc,"Verification Agent 自动生成并冻结 testbench；RTLScout 只生成和演化 RTL；固定工具链负责 compile/lint/simulation/mutation，OpenROAD 负责后端 PPA。自动测试不等于人工测试，证据归属字段也不要求人工点击。")
    add_doc_figure(doc,4,"RTLScout 内部候选演化与双 Agent 权责分离。")
    rtl_rows=[]
    for row in sorted(RTL["design_rows"],key=lambda x:x["design"]):
        m=row["mutation"]; rtl_rows.append((row["design"],row["candidate_files"],row["rtl_revisions"],f'{m["killed_count"]}/{m["executable_count"]}',f'{m["generated_count"]}/{m["invalid_count"]}',f'{m["mutation_score"]:.2f}',row["gds"]["size_bytes"]))
    add_doc_table(doc,("设计","候选","修订","杀死/可执行","生成/无效","mutation","GDS字节"),rtl_rows)
    add_doc_paragraph(doc,"边界：四题各一个生成 seed；ibex_alu 不是完整 Ibex；mutation 只评价受限 mutant，不替代形式化证明。")
    doc.add_heading("3. Agent 架构",1)
    add_doc_paragraph(doc,"八阶段为 map→semantic→experiment→hypothesis→implement→validate→review→memory。hypothesis 不可执行，Runtime 独占进程和结果登记权。四设计真实 trace 均出现八阶段；另有中断恢复、幂等和权限边界测试。")
    add_doc_table(doc,("设计","终态","事件数"),[(x["design"],x["status"],x["event_count"]) for x in AGENT["design_rows"]])
    doc.add_heading("4. EDA-to-AI 数据接口",1)
    add_doc_figure(doc,3,"原始证据、typed IR、关系表和有界查询。")
    add_doc_table(doc,("对象","数量"),[(k,v) for k,v in EDAIR["totals"].items()])
    add_doc_paragraph(doc,"边界：该消融证明比 KPI 摘要保留更多可查询结构和来源，不证明 Agent 诊断更准或 QoR 更好；当前没有 EDATracer 式大规模向量双索引。")
    doc.add_heading("5. BO/GP 参数探索",1)
    add_doc_figure(doc,6,"固定 RBF GP、重复点噪声、加权 EI、经验可行率和三轮停滞。")
    add_doc_paragraph(doc,"BO 144 次与 Random 144 次 full-flow。0.5% 阈值事件为 7/12 对 4/12；设计级中位数胜负 2:2。可写“该预算下提高阈值命中频次”，不可写普遍或统计显著优势。")
    add_doc_table(doc,("设计","BO中位utility","Random中位utility","胜者"),[(x["design"],f'{x["bo_best_utility"]["median"]*100:.3f}%',f'{x["random_best_utility"]["median"]*100:.3f}%',x["median_winner"]) for x in PARAM["design_rows"]])
    doc.add_heading("6. 自演化学习",1)
    add_doc_figure(doc,2,"事实、局部2×2干预和跨设计 holdout。")
    add_doc_figure(doc,5,"知识卡 schema 与真实 negative-transfer 样例。")
    add_doc_paragraph(doc,f"GCD 局部交互 {LEARNING['source_interaction']:.3f} µm²，FIFO holdout {LEARNING['holdout_interaction']:.3f} µm²，方向反转，系统记录 refuted 且 action_eligible=false。它证明拦截一次已观察错误迁移，不是普适因果规律。")
    doc.add_heading("7. 论文级实验总账",1)
    add_doc_table(doc,("模块","状态","aggregate","SHA-256"),[(x["capability"],x["status"],x["evidence_path"],x["evidence_sha256"]) for x in LEDGER["evidence_records"]])
    add_doc_paragraph(doc,"全仓检查：327 passed、1 deselected；Node、compileall、JSON、diff check 与 tracked-secret scan 通过。不同 aggregate 可能复用底层 run，不能把数字直接相加。")
    doc.add_heading("8. 分工与后续工作包",1)
    add_doc_table(doc,("负责人","当前职责","后续交付"),[("袁文杰","平台/Agent/EDAIR/学习/BO/实验","扩大独立设计单元与论文消融"),("刘宏博","RTL生成与验证","多seed、层级SpecIR、coverage/formal"),("沈夏南","TaiWei3D","设计×工艺矩阵和3D Agent"),("戴家乐","有限动作工具","TimingECO/Resynth/EvoDRC等逐个受保护接入")])
    doc.add_heading("9. 关键参考与对标",1)
    refs=[]
    for group in PAPERS.values():
        for title,venue,identifier,method,mapping in group:
            refs.append((title,venue,identifier,method,mapping))
    add_doc_table(doc,("研究","来源","标识","方法","平台对齐/差距"),refs)
    doc.add_heading("10. 结论",1)
    add_doc_paragraph(doc,"v2 已形成可执行、可恢复、可审计的参数级自主闭环，并有真实多设计 OpenROAD 证据。它已经越过 toy-show 的“一次 LLM 调用+单次最好值”，但仍须扩大 RTL generation seeds、独立设计单元、EDAIR诊断题和因果 holdout，才能支撑更强论文结论。有限 repair 工具按下一工作包逐个接入，不提前污染当前主线。")
    footer=doc.sections[0].footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; footer.add_run("OpenROAD AgenticEDA v2 · 2026-08-25 · Evidence-bounded report")
    doc.core_properties.title="OpenROAD AgenticEDA v2 验收与开题报告"; doc.core_properties.author="项目组"
    doc.save(MASTER_DOCX)


def main() -> int:
    DESKTOP.mkdir(parents=True, exist_ok=True)
    MASTER_HTML.write_text(master_html(),encoding="utf-8")
    modules = [
        ("01_RTL生成与验证_真实证据和前沿对标.html","RTL生成与验证",rtl_section()),
        ("02_Agent架构_真实证据和前沿对标.html","Agent架构",agent_section()),
        ("03_EDA到AI数据接口_真实证据和前沿对标.html","EDA-to-AI数据接口",edair_section()),
        ("04_BO_GP参数探索_真实证据和前沿对标.html","BO/GP参数探索",bo_section()),
        ("05_自演化学习_真实证据和前沿对标.html","自演化学习",learning_section()),
        ("06_论文级实验清单与边界.html","论文级实验清单",experiments_section()),
    ]
    for filename,title,section in modules:
        (DESKTOP/filename).write_text(module_html(title,section),encoding="utf-8")
    build_docx()
    print(MASTER_HTML); print(MASTER_DOCX)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
